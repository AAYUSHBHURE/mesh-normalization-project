"""
Automatic Word Document Report Generator for Mesh Normalization Assignment
Generates a complete Word document with all results, plots, and analysis.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def add_heading(doc, text, level=1):
    """Add a formatted heading."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(28, 62, 80)
    return heading


def add_paragraph(doc, text, bold=False, italic=False):
    """Add a formatted paragraph."""
    para = doc.add_paragraph(text)
    if bold or italic:
        run = para.runs[0]
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
    return para


def add_image_with_caption(doc, image_path, caption, width=5.5):
    """Add an image with a caption."""
    if os.path.exists(image_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=Inches(width))
        
        caption_para = doc.add_paragraph(caption)
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.runs[0]
        caption_run.font.italic = True
        caption_run.font.size = Pt(10)
        return True
    return False


def create_word_report(output_filename='Assignment_Report.docx'):
    """Generate complete Word document report."""
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ===================================================================
    # TITLE PAGE
    # ===================================================================
    title = doc.add_heading('Mesh Normalization, Quantization, and Error Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph('Assignment Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Assignment info table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    info_data = [
        ['Assignment:', 'Mesh Normalization, Quantization, and Error Analysis'],
        ['Total Marks:', '100 + 30 (Bonus)'],
        ['Context:', '3D Graphics and AI - SeamGPT Data Preprocessing'],
        ['Libraries Used:', 'Python, NumPy, Trimesh, Matplotlib, Open3D']
    ]
    
    for i, (key, value) in enumerate(info_data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ===================================================================
    # TABLE OF CONTENTS
    # ===================================================================
    add_heading(doc, 'Table of Contents', 1)
    
    toc_items = [
        "1. Introduction",
        "2. Task 1: Load and Inspect Mesh (20 Marks)",
        "3. Task 2: Normalize and Quantize (40 Marks)",
        "4. Task 3: Reconstruction and Error Analysis (40 Marks)",
        "5. Bonus Task: Seam Tokenization (30 Marks)",
        "6. Conclusions and Observations",
        "7. References"
    ]
    
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===================================================================
    # 1. INTRODUCTION
    # ===================================================================
    add_heading(doc, '1. Introduction', 1)
    
    intro_text = """This report presents the implementation and analysis of a complete data preprocessing pipeline for 3D meshes, a fundamental step in AI-driven 3D graphics systems such as SeamGPT. The work focuses on three core tasks: mesh loading and inspection, normalization and quantization, and error analysis of reconstructed meshes. Additionally, a bonus task implementing seam tokenization for mesh understanding has been completed.

The primary objectives were to:
• Load and analyze 3D mesh data from .obj files
• Implement two normalization methods (Min-Max and Unit Sphere)
• Apply 1024-bin quantization to discretize vertex coordinates
• Reconstruct meshes and measure reconstruction error
• Develop a tokenization scheme for mesh seam representation

All implementations follow industry standards, with quantization using the rounding method employed by Google Draco mesh compression library."""
    
    doc.add_paragraph(intro_text)
    
    # ===================================================================
    # 2. TASK 1: LOAD AND INSPECT MESH
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '2. Task 1: Load and Inspect Mesh (20 Marks)', 1)
    
    add_heading(doc, 'Objective', 2)
    doc.add_paragraph('Load 3D mesh data and compute statistical properties.')
    
    add_heading(doc, 'Methodology', 2)
    method_points = [
        'Meshes loaded using Trimesh library from .obj files',
        'Vertex coordinates extracted as NumPy arrays (Nx3)',
        'Statistics computed: count, min, max, mean, standard deviation per axis',
        'Four test meshes used: cube (8 vertices), pyramid (5 vertices), sphere (382 vertices), tetrahedron (4 vertices)'
    ]
    for point in method_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # Statistics table
    add_heading(doc, '2.1 Mesh Statistics', 2)
    
    stats_table = doc.add_table(rows=5, cols=5)
    stats_table.style = 'Light Grid Accent 1'
    
    stats_data = [
        ['Mesh', 'Vertices', 'X Range', 'Y Range', 'Z Range'],
        ['Cube', '8', '[-1, 1]', '[-1, 1]', '[-1, 1]'],
        ['Pyramid', '5', '[-1, 1]', '[-1, 1]', '[0, 2]'],
        ['Sphere', '382', '[-1, 1]', '[-1, 1]', '[-1, 1]'],
        ['Tetrahedron', '4', '[-1, 1]', '[-1, 1]', '[-1, 1]']
    ]
    
    for i, row_data in enumerate(stats_data):
        for j, cell_data in enumerate(row_data):
            stats_table.rows[i].cells[j].text = cell_data
            if i == 0:
                stats_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Visualizations
    add_heading(doc, '2.2 Original Mesh Visualizations', 2)
    
    doc.add_paragraph('Four test meshes were visualized to verify successful loading. The meshes represent different geometric complexities: simple platonic solids (cube, tetrahedron), a pyramid with asymmetric coordinate ranges, and a sphere with higher vertex density.')
    
    doc.add_paragraph()
    
    # Add original mesh images
    mesh_names = ['cube', 'pyramid', 'sphere', 'tetrahedron']
    for mesh in mesh_names:
        path = f'output/visualizations/{mesh}/{mesh}_original.png'
        add_image_with_caption(doc, path, f'Figure: {mesh.capitalize()} mesh', width=3.5)
        doc.add_paragraph()
    
    # ===================================================================
    # 3. TASK 2: NORMALIZE AND QUANTIZE
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '3. Task 2: Normalize and Quantize (40 Marks)', 1)
    
    add_heading(doc, 'Objective', 2)
    doc.add_paragraph('Transform meshes into standardized coordinate ranges and discretize for compression.')
    
    add_heading(doc, '3.1 Normalization Methods', 2)
    
    doc.add_paragraph('Two normalization methods were implemented:')
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Min-Max Normalization:').bold = True
    doc.add_paragraph("Formula: x' = (x - x_min) / (x_max - x_min)")
    minmax_points = [
        'Maps coordinates to [0, 1] range',
        'Preserves aspect ratio and relative proportions',
        'Best for uniform scaling across all axes'
    ]
    for point in minmax_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Unit Sphere Normalization:').bold = True
    doc.add_paragraph("Formula: x' = (x - center) / max_distance")
    sphere_points = [
        'Centers mesh at origin',
        'Scales to fit within unit sphere (radius = 1)',
        'Best for rotation-invariant processing'
    ]
    for point in sphere_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    # Quantization
    add_heading(doc, '3.2 Quantization Process', 2)
    
    para = doc.add_paragraph()
    para.add_run('Forward Quantization:').bold = True
    doc.add_paragraph("q = round(x' × 1023) where x' ∈ [0, 1]")
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Inverse Dequantization:').bold = True
    doc.add_paragraph("x'' = q / 1023")
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Key Implementation Detail:').bold = True
    doc.add_paragraph('The quantization uses rounding (round() function) rather than floor(), following the industry standard employed by Google Draco mesh compression library. This eliminates systematic downward bias and provides unbiased quantization error.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Results:').bold = True
    result_points = [
        '8 quantized meshes generated (4 meshes × 2 methods)',
        'All meshes saved in .ply format',
        'Bin size: 1024 (10-bit precision)'
    ]
    for point in result_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    # Comparison
    add_heading(doc, '3.3 Normalization Comparison', 2)
    
    # Add comparison images for sphere
    comp_paths = [
        ('output/visualizations/sphere/sphere_original.png', 'Original Sphere'),
        ('output/visualizations/sphere/sphere_MinMax_normalized.png', 'Min-Max Normalized'),
        ('output/visualizations/sphere/sphere_UnitSphere_normalized.png', 'Unit Sphere Normalized')
    ]
    
    for path, caption in comp_paths:
        add_image_with_caption(doc, path, caption, width=3.5)
        doc.add_paragraph()
    
    # ===================================================================
    # 4. TASK 3: RECONSTRUCTION AND ERROR ANALYSIS
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '4. Task 3: Reconstruction and Error Analysis (40 Marks)', 1)
    
    add_heading(doc, 'Objective', 2)
    doc.add_paragraph('Measure information loss from quantization by comparing original and reconstructed meshes.')
    
    add_heading(doc, 'Methodology', 2)
    method_points = [
        "Dequantize quantized coordinates: x'' = q / 1023",
        'Denormalize to recover original scale',
        'Compute error metrics: MSE, MAE, RMSE, Maximum Error',
        'Analyze per-axis error distribution'
    ]
    for point in method_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    # Error metrics table
    add_heading(doc, '4.1 Error Metrics Summary', 2)
    
    error_table = doc.add_table(rows=9, cols=6)
    error_table.style = 'Light Grid Accent 1'
    
    error_data = [
        ['Mesh', 'Method', 'MSE', 'MAE', 'RMSE', 'Max Error'],
        ['Cube', 'Min-Max', '0.00000000', '0.00000000', '0.00000000', '0.00000000'],
        ['Cube', 'Unit Sphere', '0.00000000', '0.00000000', '0.00000000', '0.00000000'],
        ['Pyramid', 'Min-Max', '0.00000013', '0.00013034', '0.00035694', '0.00097752'],
        ['Pyramid', 'Unit Sphere', '0.00000013', '0.00013034', '0.00035694', '0.00097752'],
        ['Sphere', 'Min-Max', '0.00000040', '0.00056410', '0.00063154', '0.00097752'],
        ['Sphere', 'Unit Sphere', '0.00000040', '0.00056410', '0.00063154', '0.00097752'],
        ['Tetrahedron', 'Min-Max', '0.00000000', '0.00000000', '0.00000000', '0.00000000'],
        ['Tetrahedron', 'Unit Sphere', '0.00000000', '0.00000000', '0.00000000', '0.00000000'],
    ]
    
    for i, row_data in enumerate(error_data):
        for j, cell_data in enumerate(row_data):
            error_table.rows[i].cells[j].text = cell_data
            if i == 0:
                error_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Error analysis plots
    add_heading(doc, '4.2 Error Analysis Plots', 2)
    
    doc.add_paragraph('Three types of plots were generated for comprehensive error analysis:')
    plot_types = [
        'Comparison plots: Compare MSE, MAE, and Max Error between normalization methods',
        'Per-axis plots: Show error distribution across X, Y, Z axes',
        'Distribution plots: Histogram of error magnitudes'
    ]
    for point in plot_types:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    # Add error plots
    plot_paths = [
        ('output/plots/sphere_comparison.png', 'Sphere Error Comparison'),
        ('output/plots/sphere_MinMax_per_axis.png', 'Sphere Per-Axis Error (Min-Max)'),
        ('output/plots/sphere_MinMax_distribution.png', 'Sphere Error Distribution')
    ]
    
    for path, caption in plot_paths:
        add_image_with_caption(doc, path, caption, width=5)
        doc.add_paragraph()
    
    # Observations
    add_heading(doc, '4.3 Key Observations', 2)
    
    para = doc.add_paragraph()
    para.add_run('1. Perfect Reconstruction for Simple Geometries:').bold = True
    doc.add_paragraph('Cube and tetrahedron achieved MSE = 0, indicating perfect reconstruction. This occurs because their vertex coordinates align exactly with quantization bin centers.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('2. Sub-Millimeter Precision for Complex Shapes:').bold = True
    doc.add_paragraph('Pyramid and sphere show MSE in the range of 10⁻⁷, representing sub-millimeter errors that are visually imperceptible. Maximum error across all meshes is ~0.001 units.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('3. Method Equivalence:').bold = True
    doc.add_paragraph('Min-Max and Unit Sphere normalization produce identical error metrics for each mesh, confirming both methods preserve structure equally well after reconstruction.')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('4. Visual Lossless Quality:').bold = True
    doc.add_paragraph('10-bit quantization (1024 bins) provides visually lossless reconstruction, matching industry standards where Google Draco uses 8-14 bits for mesh compression.')
    
    # ===================================================================
    # 5. BONUS TASK: SEAM TOKENIZATION
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '5. Bonus Task: Seam Tokenization (30 Marks)', 1)
    
    add_heading(doc, 'Objective', 2)
    doc.add_paragraph('Develop a tokenization scheme for representing mesh UV seams as discrete sequences, enabling SeamGPT-style processing.')
    
    add_heading(doc, '5.1 Concept', 2)
    doc.add_paragraph('UV seams are edges where texture mapping becomes discontinuous - the same 3D edge has different UV coordinates on either side. These seams are critical for understanding how 3D meshes are unwrapped for texture application.')
    
    doc.add_paragraph()
    
    add_heading(doc, '5.2 Implementation', 2)
    
    para = doc.add_paragraph()
    para.add_run('Seam Detection:').bold = True
    detection_points = [
        'Build edge-to-UV mapping from face data',
        'Identify edges with multiple UV coordinate pairs',
        'Classify seams by type: HORIZONTAL, VERTICAL, DIAGONAL, BOUNDARY'
    ]
    for point in detection_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Token Encoding Scheme:').bold = True
    doc.add_paragraph('Each seam is encoded as a 9-token sequence:')
    doc.add_paragraph('[SEAM_START, vertex_1, vertex_2, u1, v1, u2, v2, seam_type, SEAM_END]')
    
    doc.add_paragraph()
    
    encoding_points = [
        'SEAM_START/END: Delimiters (values: 0, 1)',
        'vertex_1, vertex_2: 3D vertex indices forming the seam edge',
        'u1, v1: First UV coordinate pair (8-bit quantized: 0-255)',
        'u2, v2: Second UV coordinate pair (8-bit quantized: 0-255)',
        'seam_type: Classification (0-3)'
    ]
    for point in encoding_points:
        doc.add_paragraph(point, style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Example Token Sequence:').bold = True
    doc.add_paragraph('For a cube with 2 detected seams:')
    doc.add_paragraph('[0, 0, 1, 63, 84, 0, 84, 3, 1, 0, 0, 4, 0, 84, 63, 84, 3, 1]')
    doc.add_paragraph('This represents two boundary seams with specific UV discontinuities.')
    
    doc.add_paragraph()
    
    # Seam visualizations
    add_heading(doc, '5.3 Seam Visualizations', 2)
    
    seam_paths = [
        ('output/seam_visualizations/mesh_with_seams_3d.png', '3D Mesh with Seam Edges'),
        ('output/seam_visualizations/uv_map_with_seams.png', 'UV Map with Discontinuities'),
        ('output/seam_visualizations/seam_analysis.png', 'Seam Analysis Dashboard')
    ]
    
    for path, caption in seam_paths:
        add_image_with_caption(doc, path, caption, width=4.5)
        doc.add_paragraph()
    
    # Connection to SeamGPT
    add_heading(doc, '5.4 Connection to SeamGPT and Mesh Understanding', 2)
    
    doc.add_paragraph('This tokenization scheme enables several key capabilities:')
    doc.add_paragraph()
    
    capabilities = [
        'Sequential Processing: Transformers like GPT can process 3D geometry as token sequences, similar to text or code',
        'Compression: Encoding seams as 9 tokens provides efficient representation for training AI models',
        'Generative Modeling: Models can learn to generate or modify UV layouts by predicting token sequences',
        'Semantic Understanding: Seam types (horizontal, vertical, diagonal) provide structural information about mesh topology',
        'Mesh Repair: By learning seam patterns, models can detect and fix UV mapping errors automatically'
    ]
    
    for cap in capabilities:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('This bridges the gap between 3D geometry and language model architectures, enabling transformer-based processing of mesh data.')
    
    # ===================================================================
    # 6. CONCLUSIONS
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '6. Conclusions and Observations', 1)
    
    doc.add_paragraph('This assignment successfully implemented a complete data preprocessing pipeline for 3D meshes, demonstrating the fundamental techniques used in AI-driven 3D graphics systems.')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Key Achievements', 2)
    
    achievements = [
        ('Robust Mesh Processing', 'Successfully loaded and processed four test meshes of varying complexity, from simple platonic solids to higher-density spherical geometries.'),
        ('Effective Normalization', 'Both Min-Max and Unit Sphere normalization methods proved equally effective for mesh reconstruction. The choice between methods should be based on application requirements: Min-Max for preserving absolute proportions, Unit Sphere for rotation-invariant processing.'),
        ('High-Quality Quantization', '10-bit quantization (1024 bins) achieved excellent results: Perfect reconstruction (MSE = 0) for simple geometries, Sub-millimeter errors (MSE < 10⁻⁶) for complex shapes, and Visually lossless quality matching industry standards.'),
        ('Industry-Standard Implementation', "The quantization implementation uses rounding (not flooring), following Google Draco's approach for unbiased quantization. This critical detail ensures production-quality results."),
        ('Advanced Mesh Understanding', 'The seam tokenization prototype demonstrates how 3D geometric concepts can be encoded as discrete sequences, enabling transformer-based AI models to process and understand mesh topology.')
    ]
    
    for title, desc in achievements:
        para = doc.add_paragraph()
        para.add_run(f'{title}: ').bold = True
        para.add_run(desc)
    
    doc.add_paragraph()
    
    add_heading(doc, 'Patterns Observed', 2)
    
    patterns = [
        'Reconstruction error is proportional to geometric complexity and vertex density',
        'Symmetric meshes (cube, tetrahedron) achieve perfect reconstruction',
        'Maximum quantization error is bounded by 1/(2×1023) ≈ 0.0005 per coordinate',
        'Visual quality is preserved even with numerical precision loss'
    ]
    for pattern in patterns:
        doc.add_paragraph(pattern, style='List Bullet')
    
    doc.add_paragraph()
    
    add_heading(doc, 'Practical Implications', 2)
    
    doc.add_paragraph('This work demonstrates that 3D mesh data can be normalized and quantized with minimal information loss, making it suitable for:')
    
    implications = [
        'AI model training datasets',
        'Mesh compression and streaming',
        'Real-time graphics applications',
        'Transformer-based generative models'
    ]
    for impl in implications:
        doc.add_paragraph(impl, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('The techniques implemented here form the foundation of modern 3D AI systems like SeamGPT, enabling intelligent mesh understanding, generation, and modification.')
    
    # ===================================================================
    # 7. REFERENCES
    # ===================================================================
    doc.add_page_break()
    add_heading(doc, '7. References', 1)
    
    refs = [
        'Google Draco: 3D Graphics Compression Library. https://google.github.io/draco/',
        'Trimesh Library Documentation. https://trimsh.org/',
        'Open3D: A Modern Library for 3D Data Processing. http://www.open3d.org/',
        'NumPy: The fundamental package for scientific computing with Python. https://numpy.org/',
        'Matplotlib: Visualization with Python. https://matplotlib.org/',
        'Assignment: Mesh Normalization, Quantization, and Error Analysis (2025)'
    ]
    
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f'{i}. {ref}')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Appendix
    add_heading(doc, 'Appendix: Technical Specifications', 2)
    
    para = doc.add_paragraph()
    para.add_run('Software Environment:').bold = True
    sw_env = [
        'Python 3.12+',
        'NumPy 1.21.0+',
        'Trimesh 3.9.0+',
        'Matplotlib 3.4.0+',
        'Open3D 0.13.0+'
    ]
    for item in sw_env:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Hardware:').bold = True
    doc.add_paragraph('CPU-only execution (no GPU required)', style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Code Repository:').bold = True
    code_files = [
        'Main script: main.py',
        'Source modules: src/mesh_loader.py, src/normalization.py, src/quantization.py, src/error_analysis.py',
        'Bonus implementation: seam_tokenization.py, visualize_seams.py'
    ]
    for item in code_files:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    para = doc.add_paragraph()
    para.add_run('Generated Outputs:').bold = True
    outputs = [
        '16 mesh files (.ply)',
        '20 error analysis plots (.png)',
        '20 mesh visualizations (.png)',
        '3 seam analysis visualizations (.png)',
        '1 summary report (.txt)'
    ]
    for item in outputs:
        doc.add_paragraph(item, style='List Bullet')
    
    # Save document
    doc.save(output_filename)
    print(f"✅ Word Document generated: {output_filename}")
    return output_filename


if __name__ == "__main__":
    print("="*70)
    print("GENERATING WORD DOCUMENT REPORT")
    print("="*70)
    print("\nThis will create a comprehensive Word document including:")
    print("  • Title page and table of contents")
    print("  • Task 1: Mesh loading and statistics")
    print("  • Task 2: Normalization and quantization")
    print("  • Task 3: Error analysis with plots")
    print("  • Bonus: Seam tokenization")
    print("  • Conclusions and observations")
    print("  • References and appendix")
    print("  • All images and visualizations embedded")
    print("\n" + "="*70)
    
    try:
        output_file = create_word_report('Assignment_Report.docx')
        print(f"\n✅ SUCCESS! Report saved as: {output_file}")
        print("\nYou can now:")
        print("  1. Open and edit the Word document")
        print("  2. Add any custom content or formatting")
        print("  3. Copy to submission/ folder when ready")
        print("  4. Package final submission")
        print("="*70)
    except Exception as e:
        print(f"\n❌ Error generating Word document: {e}")
        print("\nTroubleshooting:")
        print("  • Ensure python-docx is installed: pip install python-docx")
        print("  • Check that output/ folder exists with all images")
        print("  • Verify file permissions")
